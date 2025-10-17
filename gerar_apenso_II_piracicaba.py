# -*- coding: utf-8 -*-
"""
Script: gerar_apenso_II_piracicaba.py
Autor: Rafael Neves
Descrição:
Gera o arquivo "Apenso II - PIRACICABA (Final e Formatado).docx"
a partir dos arquivos:
- Apenso II original para preenchimento.docx
- Ordem de visita tecnica.pdf
- Visita Piracicaba com valores para mandar.pdf
"""

import fitz  # PyMuPDF
from docx import Document


# === 1️⃣ Leitura do PDF de visita técnica ===
def extrair_dados_visita_tecnica(pdf_path):
    dados = []
    with fitz.open(pdf_path) as pdf:
        texto = ""
        for pagina in pdf:
            texto += pagina.get_text("text")

    linhas = texto.split("\n")
    for linha in linhas:
        if any(ch.isdigit() for ch in linha) and "BTUS" in linha.upper():
            partes = linha.split()
            patrimonio = [p for p in partes if "-" in p or p.isdigit()]
            patrimonio = patrimonio[0] if patrimonio else "N/D"
            btus = [p for p in partes if "BTUS" in p.upper()]
            btus = btus[0].replace("BTUS", "").replace(",", "").strip() if btus else "N/D"
            local = linha.split("MARCA")[-1].split()[-3:]
            local = " ".join(local)
            marca = "MIDEA" if "MIDEA" in linha.upper() else \
                    "ELGIN" if "ELGIN" in linha.upper() else \
                    "SAMSUNG" if "SANSUNG" in linha.upper() else \
                    "HITACHI" if "HITACHI" in linha.upper() else "N/D"
            dados.append({
                "patrimonio": patrimonio,
                "marca": marca,
                "btus": btus,
                "local": local
            })
    return dados


# === 2️⃣ Leitura da planilha PDF (Visita Piracicaba com valores) ===
def extrair_valores_planilha(pdf_path):
    dados = []
    with fitz.open(pdf_path) as pdf:
        texto = ""
        for pagina in pdf:
            texto += pagina.get_text("text")

    linhas = [l for l in texto.split("\n") if "R$" in l]
    for linha in linhas:
        partes = linha.split("R$")
        descricao = partes[0].strip()
        valor = "R$ " + partes[1].strip().split()[0] if len(partes) > 1 else "N/D"
        secao = descricao.split("-")[-1].strip() if "-" in descricao else descricao
        dados.append({
            "descricao": descricao,
            "valor": valor,
            "secao": secao
        })
    return dados


# === 3️⃣ Função principal de geração ===
def gerar_apenso_II(modelo_path, visita_pdf, planilha_pdf, saida_path):
    doc_base = Document(modelo_path)
    dados_tecnicos = extrair_dados_visita_tecnica(visita_pdf)
    dados_valores = extrair_valores_planilha(planilha_pdf)

    # Criar os novos blocos
    doc_base.add_page_break()
    doc_base.add_paragraph("NOVOS SERVIÇOS – PIRACICABA")
    doc_base.add_paragraph("="*80)

    for i, item in enumerate(dados_tecnicos, 1):
        # Encontrar a seção correspondente na planilha
        descricao, valor = "N/D", "N/D"
        for v in dados_valores:
            if v["secao"].lower().strip() in item["local"].lower().strip():
                descricao, valor = v["descricao"], v["valor"]
                break

        doc_base.add_paragraph(f"\nSERVIÇO {i}")
        doc_base.add_paragraph("DADOS DO APARELHO CONDICIONADOR DE AR")
        doc_base.add_paragraph(f"MARCA: {item['marca']}")
        doc_base.add_paragraph("MODELO: SPLIT")
        doc_base.add_paragraph(f"PATRIMÔNIO: {item['patrimonio']}")
        doc_base.add_paragraph(f"BTUs: {item['btus']}")
        doc_base.add_paragraph(f"LOCAL ONDE ESTÁ INSTALADO: {item['local']}")
        doc_base.add_paragraph("OFICIAL GESTOR DA ATA DE REGISTRO DE PREÇOS: ")
        doc_base.add_paragraph("")
        doc_base.add_paragraph("PLANILHA DE COMPOSIÇÃO DE CUSTOS")
        doc_base.add_paragraph("Descrição dos serviços a serem realizados")
        doc_base.add_paragraph(descricao)
        doc_base.add_paragraph("Quantidade: 1")
        doc_base.add_paragraph(f"Valor Unitário: {valor}")
        doc_base.add_paragraph(f"Valor Total: {valor}")
        doc_base.add_paragraph(f"VALOR TOTAL DOS SERVIÇOS: {valor}")
        doc_base.add_paragraph("="*80)

    doc_base.save(saida_path)
    print(f"✅ Arquivo gerado com sucesso: {saida_path}")


# === 4️⃣ EXECUÇÃO ===
if __name__ == "__main__":
    gerar_apenso_II(
        modelo_path="Apenso II original para preenchimento.docx",
        visita_pdf="Ordem de visita tecnica.pdf",
        planilha_pdf="Visita Piracicaba com valores para mandar.pdf",
        saida_path="Apenso II - PIRACICABA (Final e Formatado).docx"
    )

